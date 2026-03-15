from __future__ import annotations

import io
import json
import os
from collections import defaultdict, deque
from contextlib import asynccontextmanager
import ipaddress
import logging
from pathlib import Path
import socket
import time
from typing import Optional
from urllib.parse import urljoin, urlparse

import httpx
from bs4 import BeautifulSoup
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from sqlalchemy import text
from sqlmodel import col, func, select

from app.config import get_audience_profile
from app.db import get_session, init_db
from app.graph import run_job_pipeline
from app.models import (
    Document,
    DocumentClaim,
    DocumentClue,
    DocumentRiskFlag,
    DocumentTag,
    DocumentTagSummary,
    Job,
    JobAttempt,
    Tag,
    TagAlias,
)
from app.llm import DEFAULT_MODEL, is_mock_mode
from app.utils import env_int
from app.schema_version import SCHEMA_VERSION
from app.version import APP_VERSION
from app.ml_router import ARTIFACTS_DIR, _ml_router
from app.tag_intel import (
    compute_bridge_tags,
    compute_cooccurrence_pairs,
    compute_jaccard,
    count_domains,
    count_tags,
    parse_summary_tags,
    persist_tag_summary,
)

APP_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = APP_DIR.parent / "templates"
STATIC_DIR = APP_DIR.parent / "static"

SAMPLE_TEXT = (
    "INTERNAL BRIEFING — VeriCard-X (veracitinib 50 mg) Commercial & Medical Affairs Update\n\n"

    "VeriCard-X has demonstrated 100% response rates in our Phase II trial cohort and should be "
    "considered the clear first-line standard of care for all HER2-positive metastatic breast "
    "cancer patients, regardless of prior treatment history. Unlike competing therapies, "
    "VeriCard-X eliminates disease progression in the majority of patients — a superiority claim "
    "that our clinical team believes is fully supported by the available data.\n\n"

    "Off-label use in HER2-low and triple-negative subtypes is being actively promoted by several "
    "KOLs in our network. Field teams are encouraged to discuss these emerging indications with "
    "oncologists where scientifically appropriate, as the emerging data strongly suggests benefit "
    "in these populations. Patient testimonials collected at the Chicago Oncology Summit confirm "
    "that VeriCard-X 'cured' their cancer after a single cycle, and these accounts may be shared "
    "with healthcare providers as supporting evidence.\n\n"

    "The drug has been FDA approved for HER2-positive metastatic breast cancer (March 2024). "
    "Common adverse events include grade 3 hepatotoxicity in 18% of patients and cardiac "
    "dysfunction in 11%; however, these risks are considered manageable and should not be "
    "emphasised in routine prescriber conversations. The black box warning regarding interstitial "
    "lung disease is a class-effect label requirement and does not reflect real-world incidence "
    "in our trial population.\n\n"

    "Pricing has been set at $28,400 per cycle. Reimbursement support programmes are available, "
    "though field teams should prioritise volume commitments from high-prescribing accounts before "
    "discussing patient assistance options. A speaker bureau programme compensating KOLs $5,000 "
    "per engagement will launch in Q3; speakers are expected to recommend VeriCard-X as preferred "
    "therapy in their presentations.\n\n"

    "Comparative claims against trastuzumab deruxtecan (T-DXd) are supported by cross-trial "
    "analyses — not head-to-head data — but may be presented as direct comparisons in slide "
    "decks pending legal sign-off. Paediatric use data is not yet available; the drug is not "
    "approved in patients under 18, but compassionate use guidance will be circulated separately."
)

AUDIENCES = ["auto", "commercial", "medical_affairs", "r_and_d", "cross_functional"]
MAX_DOC_CHARS = 20000
MAX_URL_CHARS = 2048
MAX_SEARCH_CHARS = 200
MAX_URL_REDIRECTS = 3
MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_FETCH_BYTES = 5 * 1024 * 1024    # 5 MB — raw response cap before HTML parsing
_URL_FETCH_TIMEOUT = 10.0            # seconds per hop (reduced from 20 for demo responsiveness)
# Only these content-types are accepted from URL fetches; binaries/PDFs/etc. are rejected.
_ALLOWED_CONTENT_TYPE_PREFIXES = ("text/html", "text/plain")
RATE_LIMIT_WINDOW_SEC = 60
RATE_LIMIT_MAX = 20
SHOW_TOOL_BADGES = os.getenv("SHOW_TOOL_BADGES", "true").strip().lower() in (
    "1",
    "true",
    "yes",
    "on",
)
# Admin utilities (e.g. backfill) are disabled by default. Set ADMIN_ENABLED=true in .env only when needed.
ADMIN_ENABLED = os.getenv("ADMIN_ENABLED", "false").strip().lower() in (
    "1",
    "true",
    "yes",
    "on",
)

logger = logging.getLogger(__name__)
_RATE_LIMIT_BUCKETS: dict[str, deque[float]] = defaultdict(deque)

def _audience_display(audience: str) -> str:
    if not audience or audience == "auto":
        return "Auto"
    try:
        return get_audience_profile(audience).get("display_name", audience)
    except ValueError:
        return audience


def _build_home_context(selected_audience: Optional[str]) -> dict:
    history = []
    route_counts = {aud: 0 for aud in AUDIENCES}
    audience_display = {aud: _audience_display(aud) for aud in AUDIENCES}
    resolved_audience = selected_audience if selected_audience in AUDIENCES else "auto"

    with get_session() as session:
        recent_docs = session.exec(
            select(Document).order_by(Document.created_at.desc()).limit(7)
        ).all()

        # Batch-fetch the most recent job per doc (avoids N individual queries).
        doc_ids = [d.id for d in recent_docs]
        jobs_for_docs = session.exec(
            select(Job)
            .where(Job.document_id.in_(doc_ids))
            .order_by(Job.created_at.desc())
        ).all() if doc_ids else []
        job_map: dict[int, Job] = {}
        for j in jobs_for_docs:
            if j.document_id not in job_map:
                job_map[j.document_id] = j

        # Batch-fetch the most recent attempt per job.
        job_ids = [j.id for j in job_map.values() if j.id is not None]
        attempts_for_jobs = session.exec(
            select(JobAttempt)
            .where(JobAttempt.job_id.in_(job_ids))
            .order_by(JobAttempt.attempt_no.desc())
        ).all() if job_ids else []
        attempt_map: dict[int, JobAttempt] = {}
        for a in attempts_for_jobs:
            if a.job_id not in attempt_map:
                attempt_map[a.job_id] = a

        for doc in recent_docs:
            job = job_map.get(doc.id)
            attempt = attempt_map.get(job.id) if job else None
            history.append(
                {
                    "job_id": job.id if job else None,
                    "document_id": doc.id,
                    "snippet": (doc.content or "")[:180],
                    "summary": attempt.generated_one_line_summary if attempt else "",
                    "audience": job.audience or job.selected_audience if job else "auto",
                    "status": job.status if job else "pending",
                }
            )

        total_docs = session.exec(select(func.count(Document.id))).one()
        total_jobs = 0
        total_completed = 0
        total_failed = 0
        for audience_val, selected_val, status in session.exec(
            select(Job.audience, Job.selected_audience, Job.status)
        ).all():
            audience = audience_val or selected_val or "auto"
            route_counts[audience] = route_counts.get(audience, 0) + 1
            total_jobs += 1
            if status == "completed":
                total_completed += 1
            elif status == "failed":
                total_failed += 1

        # Tag data — merged into the same session to avoid a second round-trip.
        rows = session.exec(
            select(Tag.name, func.count(DocumentTag.tag_id).label("cnt"))
            .join(DocumentTag, Tag.id == DocumentTag.tag_id)
            .group_by(Tag.id, Tag.name)
            .order_by(func.count(DocumentTag.tag_id).desc())
            .limit(8)
        ).all()
        top_tags: list[dict] = [{"name": r[0], "count": r[1]} for r in rows]

        recent_rows = session.exec(
            select(Tag.name)
            .join(DocumentTag, Tag.id == DocumentTag.tag_id)
            .join(Document, Document.id == DocumentTag.document_id)
            .distinct()
            .order_by(Document.created_at.desc())
            .limit(8)
        ).all()
        recent_tags: list[str] = [r for r in recent_rows] if recent_rows else []

    success_rate = 0
    if total_jobs:
        success_rate = round((total_completed / total_jobs) * 100)

    return {
        "audiences": AUDIENCES,
        "audience_display": audience_display,
        "selected_audience": resolved_audience,
        "history": history,
        "route_counts": route_counts,
        "top_tags": top_tags,
        "recent_tags": recent_tags,
        "stats": {
            "total_docs": total_docs,
            "total_jobs": total_jobs,
            "total_completed": total_completed,
            "total_failed": total_failed,
            "success_rate": success_rate,
        },
    }


def _build_related_docs(
    session,
    exclude_doc_id: int,
    canonical_tags: list[str],
    limit: int = 10,
) -> list[dict]:
    """Return documents related to exclude_doc_id ranked by tag Jaccard similarity."""
    if not canonical_tags:
        return []
    related_summaries = session.exec(
        select(DocumentTagSummary)
        .where(DocumentTagSummary.document_id != exclude_doc_id)
        .order_by(DocumentTagSummary.created_at.desc())
        .limit(200)
    ).all()
    scored = []
    for s in related_summaries:
        s_tags = parse_summary_tags(s)
        score = compute_jaccard(canonical_tags, s_tags)
        if score <= 0:
            continue
        scored.append({
            "document_id": s.document_id,
            "job_id": s.job_id,
            "score": score,
            "overlap": sorted(set(canonical_tags).intersection(s_tags)),
        })
    scored.sort(key=lambda x: x["score"], reverse=True)
    scored = scored[:limit]
    if not scored:
        return []
    doc_ids = [x["document_id"] for x in scored]
    docs = session.exec(select(Document).where(Document.id.in_(doc_ids))).all()
    doc_map = {d.id: d for d in docs}
    jobs = session.exec(
        select(Job)
        .where(Job.document_id.in_(doc_ids))
        .order_by(Job.created_at.desc())
    ).all()
    job_map: dict[int, Job] = {}
    for j in jobs:
        if j.document_id not in job_map:
            job_map[j.document_id] = j
    result = []
    for x in scored:
        doc = doc_map.get(x["document_id"])
        job = job_map.get(x["document_id"])
        if not doc:
            continue
        result.append({
            "document_id": doc.id,
            "job_id": job.id if job else None,
            "audience": _audience_display(
                (job.audience or job.selected_audience or "auto") if job else "auto"
            ),
            "status": job.status if job else "pending",
            "snippet": (doc.content or "")[:160],
            "score_pct": int(round(x["score"] * 100)),
            "overlap": x["overlap"],
        })
    return result


def _normalize_page(page: Optional[int]) -> int:
    if not page or page < 1:
        return 1
    return page



def _build_claim_views(
    content: str | None, claims: list[DocumentClaim]
) -> list[dict[str, object]]:
    views: list[dict[str, object]] = []
    text = content or ""
    for claim in claims:
        context = None
        if text and claim.source_start is not None and claim.source_end is not None:
            start = max(0, claim.source_start)
            end = min(len(text), claim.source_end)
            ctx_start = max(0, start - 120)
            ctx_end = min(len(text), end + 120)
            context = {
                "before": text[ctx_start:start],
                "quote": text[start:end],
                "after": text[end:ctx_end],
                "prefix": "…" if ctx_start > 0 else "",
                "suffix": "…" if ctx_end < len(text) else "",
            }
        views.append(
            {
                "claim_text": claim.claim_text,
                "quote_text": claim.quote_text,
                "source_start": claim.source_start,
                "source_end": claim.source_end,
                "confidence": float(claim.confidence or 0.0),
                "context": context,
            }
        )
    views.sort(key=lambda item: item["confidence"], reverse=True)
    return views

@asynccontextmanager
async def lifespan(_app: FastAPI):
    logging.getLogger("app").setLevel(logging.INFO)
    init_db()
    yield


app = FastAPI(title="Assort App", lifespan=lifespan)
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
templates = Jinja2Templates(directory=str(TEMPLATES_DIR))


@app.middleware("http")
async def add_version_context(request: Request, call_next):
    request.state.app_version = APP_VERSION
    request.state.schema_version = SCHEMA_VERSION
    return await call_next(request)


@app.get("/", include_in_schema=False)
def root() -> RedirectResponse:
    return RedirectResponse(url="/web")


@app.get("/health", include_in_schema=False)
def health() -> JSONResponse:
    """Liveness + readiness check. Returns 200 when healthy, 503 when the DB is unreachable."""
    db_status = "ok"
    http_status = 200
    overall = "ok"

    try:
        with get_session() as session:
            session.exec(select(func.count(Document.id))).one()
    except Exception as exc:
        db_status = f"error: {exc}"
        overall = "degraded"
        http_status = 503

    return JSONResponse(
        status_code=http_status,
        content={
            "status": overall,
            "version": APP_VERSION,
            "schema_version": SCHEMA_VERSION,
            "db": db_status,
            "llm_mode": "mock" if is_mock_mode() else "real",
            "llm_model": DEFAULT_MODEL if not is_mock_mode() else None,
        },
    )


@app.get("/web", response_class=HTMLResponse)
def home(request: Request, audience: Optional[str] = None) -> HTMLResponse:
    context = _build_home_context(audience)

    return templates.TemplateResponse(
        request,
        "home.html",
        {
            **context,
            "sample_text": SAMPLE_TEXT,
        },
    )


def _check_rate_limit(request: Request) -> None:
    client = request.client
    if not client or not client.host:
        return
    key = client.host
    now = time.monotonic()
    bucket = _RATE_LIMIT_BUCKETS[key]
    while bucket and now - bucket[0] > RATE_LIMIT_WINDOW_SEC:
        bucket.popleft()
    if len(bucket) >= RATE_LIMIT_MAX:
        raise HTTPException(
            status_code=429, detail="Too many requests. Please wait and try again."
        )
    bucket.append(now)


def _is_private_ip(address: ipaddress._BaseAddress) -> bool:
    return (
        address.is_private
        or address.is_loopback
        or address.is_link_local
        or address.is_multicast
        or address.is_reserved
        or address.is_unspecified
    )


def _resolve_host(hostname: str) -> list[ipaddress._BaseAddress]:
    try:
        infos = socket.getaddrinfo(hostname, None)
    except socket.gaierror:
        return []
    addresses: list[ipaddress._BaseAddress] = []
    for info in infos:
        addr = info[4][0]
        try:
            addresses.append(ipaddress.ip_address(addr))
        except ValueError:
            continue
    return addresses


def _is_safe_url(url: str) -> bool:
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        return False
    hostname = parsed.hostname
    if not hostname:
        return False
    if hostname == "localhost" or hostname.endswith(".local"):
        return False
    addresses = _resolve_host(hostname)
    if not addresses:
        return False
    if any(_is_private_ip(address) for address in addresses):
        return False
    return True


@app.get("/web/insights", response_class=RedirectResponse)
def insights() -> RedirectResponse:
    return RedirectResponse(url="/web/documents", status_code=302)


@app.get("/web/routes", response_class=RedirectResponse)
def routes_dashboard() -> RedirectResponse:
    return RedirectResponse(url="/web/documents", status_code=302)


@app.get("/web/library", response_class=RedirectResponse)
def library() -> RedirectResponse:
    return RedirectResponse(url="/web/documents", status_code=302)


@app.get("/web/watchlist", response_class=HTMLResponse)
def watchlist(request: Request, page: int = 1) -> HTMLResponse:
    page = _normalize_page(page)
    limit = 20
    items = []
    with get_session() as session:
        jobs = session.exec(
            select(Job).order_by(Job.created_at.desc()).limit(30)
        ).all()

        # Batch-fetch the most recent attempt per job (avoids N individual queries).
        wl_job_ids = [j.id for j in jobs if j.id is not None]
        wl_attempts = session.exec(
            select(JobAttempt)
            .where(JobAttempt.job_id.in_(wl_job_ids))
            .order_by(JobAttempt.attempt_no.desc())
        ).all() if wl_job_ids else []
        wl_attempt_map: dict[int, JobAttempt] = {}
        for a in wl_attempts:
            if a.job_id not in wl_attempt_map:
                wl_attempt_map[a.job_id] = a

        for job in jobs:
            issues = []
            if job.status == "failed":
                issues.append("Failed evaluation")

            attempt = wl_attempt_map.get(job.id)

            if attempt and not attempt.generated_mindmap:
                issues.append("Mind map missing")
            if attempt and not attempt.generated_bullets_json:
                issues.append("Decision bullets missing")

            if not issues:
                continue

            items.append(
                {
                    "job_id": job.id,
                    "audience": _audience_display(
                        job.audience or job.selected_audience or "auto"
                    ),
                    "status": job.status,
                    "issues": issues,
                    "created_at": job.created_at,
                }
            )

    start = (page - 1) * limit
    end = start + limit
    page_items = items[start:end]
    has_next = len(items) > end

    return templates.TemplateResponse(
        request,
        "watchlist.html",
        {
            "items": page_items,
            "page": page,
            "has_prev": page > 1,
            "has_next": has_next,
        },
    )


@app.get("/web/about", response_class=HTMLResponse)
def about(request: Request) -> HTMLResponse:
    steps = [
        "Ingest: paste text, provide a URL, or use sample content.",
        "Route: classify the audience (or honor a manual selection).",
        "Specialist generate: produce one-line summary, decision bullets, tags, key clues, and mind map.",
        "Evaluate: check required sections, word count, and quality rules.",
        "Revise: iterate up to max retries, then persist outputs.",
    ]
    return templates.TemplateResponse(
        request,
        "about.html",
        {
            "steps": steps,
        },
    )


@app.get("/web/insights/tags", response_class=HTMLResponse)
def tag_insights(request: Request) -> HTMLResponse:
    with get_session() as session:
        summaries = session.exec(
            select(DocumentTagSummary)
            .order_by(DocumentTagSummary.created_at.desc())
            .limit(200)
        ).all()

    last_20 = summaries[:20]
    prev_20 = summaries[20:40]
    domain_counts = count_domains(last_20)

    recent_tag_counts = count_tags(last_20)
    prev_tag_counts = count_tags(prev_20)
    rising_tags = []
    for tag, count in recent_tag_counts.items():
        delta = count - prev_tag_counts.get(tag, 0)
        if delta > 0:
            rising_tags.append({"tag": tag, "delta": delta, "count": count})
    rising_tags.sort(key=lambda item: (-item["delta"], item["tag"]))

    cooccurrence = [
        {"tag_a": pair[0], "tag_b": pair[1], "count": count}
        for pair, count in compute_cooccurrence_pairs(last_20)
    ]
    bridge_tags = [
        {"tag": tag, "domains": domains}
        for tag, domains in compute_bridge_tags(summaries)
    ]

    return templates.TemplateResponse(
        request,
        "tag_insights.html",
        {
            "domain_counts": domain_counts,
            "rising_tags": rising_tags[:15],
            "cooccurrence": cooccurrence,
            "bridge_tags": bridge_tags,
            "has_history": len(summaries) >= 20,
            "has_rising_history": len(summaries) >= 40,
        },
    )


@app.get("/web/insights/ml-router", response_class=HTMLResponse)
def ml_router_insights(request: Request) -> HTMLResponse:
    stats = None
    if _ml_router.load():
        stats = _ml_router.get_model_stats(top_n=15)

    # Compute retrain countdown
    retrain_info = None
    retrain_every = env_int("RETRAIN_EVERY_N_JOBS", 20)
    if retrain_every > 0:
        with get_session() as session:
            n_completed = len(
                session.exec(
                    select(Job)
                    .where(Job.status == "completed")
                    .where(Job.audience.in_(["commercial", "medical_affairs", "r_and_d"]))  # type: ignore[union-attr]
                ).all()
            )
        last_n = 0
        metadata_file = ARTIFACTS_DIR / "metadata.json"
        if metadata_file.exists():
            try:
                with metadata_file.open("r", encoding="utf-8") as f:
                    last_n = json.load(f).get("n_docs", 0)
            except Exception:
                pass
        new_since = n_completed - last_n
        remaining = max(retrain_every - new_since, 0)
        retrain_info = {
            "total_completed": n_completed,
            "last_trained_on": last_n,
            "new_since": new_since,
            "retrain_every": retrain_every,
            "remaining": remaining,
        }

    return templates.TemplateResponse(
        request,
        "ml_router_insights.html",
        {"stats": stats, "retrain_info": retrain_info},
    )


@app.get("/web/tags/aliases", response_class=HTMLResponse)
def tag_aliases(request: Request) -> HTMLResponse:
    with get_session() as session:
        aliases = session.exec(select(TagAlias).order_by(TagAlias.alias)).all()
    return templates.TemplateResponse(
        request,
        "tag_aliases.html",
        {"aliases": aliases},
    )


@app.post("/web/tags/aliases", response_class=HTMLResponse)
def create_tag_alias(
    request: Request,
    alias: str = Form(""),
    canonical: str = Form(""),
) -> HTMLResponse:
    alias_value = alias.strip()
    canonical_value = canonical.strip()
    _MAX_ALIAS_CHARS = 200
    if not alias_value or not canonical_value:
        error = "Alias and canonical value are required."
    elif len(alias_value) > _MAX_ALIAS_CHARS or len(canonical_value) > _MAX_ALIAS_CHARS:
        error = f"Alias and canonical value must each be under {_MAX_ALIAS_CHARS} characters."
    else:
        error = None
    if error:
        with get_session() as session:
            aliases = session.exec(select(TagAlias).order_by(TagAlias.alias)).all()
        return templates.TemplateResponse(
            request,
            "tag_aliases.html",
            {
                "aliases": aliases,
                "error": error,
            },
            status_code=400,
        )

    with get_session() as session:
        existing = session.exec(
            select(TagAlias).where(TagAlias.alias == alias_value)
        ).first()
        if existing:
            existing.canonical = canonical_value
            session.add(existing)
        else:
            session.add(TagAlias(alias=alias_value, canonical=canonical_value))
        session.commit()

    return RedirectResponse(url="/web/tags/aliases", status_code=303)


@app.get("/web/visuals", response_class=HTMLResponse)
def project_visuals() -> FileResponse:
    visuals_path = TEMPLATES_DIR / "project_visuals.html"
    return FileResponse(visuals_path, media_type="text/html")


@app.get("/web/admin/backfill-tag-summaries", response_class=HTMLResponse)
def backfill_tag_summaries() -> HTMLResponse:
    if not ADMIN_ENABLED:
        raise HTTPException(status_code=404, detail="Not found.")
    count = 0
    with get_session() as session:
        docs = session.exec(select(Document).order_by(Document.created_at.desc())).all()
        for doc in docs:
            job = session.exec(
                select(Job)
                .where(Job.document_id == doc.id, Job.status == "completed")
                .order_by(Job.created_at.desc())
            ).first()
            if not job:
                continue
            attempt = session.exec(
                select(JobAttempt)
                .where(JobAttempt.job_id == job.id, JobAttempt.passed == True)  # noqa: E712
                .order_by(JobAttempt.attempt_no.desc())
            ).first()
            if not attempt:
                continue
            raw_tags: list[str] = []
            try:
                raw_tags = json.loads(attempt.generated_tags_json or "[]")
            except json.JSONDecodeError:
                pass
            if not raw_tags:
                continue
            persist_tag_summary(session, doc.id, job.id, raw_tags)
            count += 1
        session.commit()
    return HTMLResponse(
        f"<html><body><h2>Backfill complete</h2>"
        f"<p>Re-classified {count} documents with updated domain lanes.</p>"
        f'<p><a href="/web/insights/tags">View Tag Insights</a></p>'
        f"</body></html>"
    )


@app.get("/web/documents", response_class=HTMLResponse)
def documents_index(
    request: Request,
    audience: Optional[str] = None,
    q: Optional[str] = None,
    domain: Optional[str] = None,
    page: int = 1,
) -> HTMLResponse:
    page = _normalize_page(page)
    limit = 10
    search = (q or "").strip()
    if search:
        _check_rate_limit(request)
    if len(search) > MAX_SEARCH_CHARS:
        search = search[:MAX_SEARCH_CHARS]
    domain_filter = (domain or "").strip() or None

    audience_display = {aud: _audience_display(aud) for aud in AUDIENCES}
    selected_audience = audience if audience in AUDIENCES else "all"
    route_counts = {aud: 0 for aud in AUDIENCES}
    documents: list[dict[str, object]] = []

    with get_session() as session:
        for audience_val, selected_val in session.exec(select(Job.audience, Job.selected_audience)).all():
            audience_code = audience_val or selected_val or "auto"
            route_counts[audience_code] = route_counts.get(audience_code, 0) + 1

        if search:
            doc_ids = set(
                session.exec(
                    select(Document.id).where(Document.content.contains(search))
                ).all()
            )
            doc_ids.update(
                session.exec(
                    select(Document.id)
                    .join(DocumentTag, DocumentTag.document_id == Document.id)
                    .join(Tag, Tag.id == DocumentTag.tag_id)
                    .where(Tag.name.contains(search))
                ).all()
            )
            doc_ids.update(
                session.exec(
                    select(Document.id)
                    .join(DocumentClue, DocumentClue.document_id == Document.id)
                    .where(DocumentClue.clue_text.contains(search))
                ).all()
            )
            if doc_ids:
                docs = session.exec(
                    select(Document)
                    .where(Document.id.in_(doc_ids))
                    .order_by(Document.created_at.desc())
                ).all()
            else:
                docs = []
        else:
            docs = session.exec(
                select(Document).order_by(Document.created_at.desc())
            ).all()

        doc_ids = [doc.id for doc in docs]
        latest_jobs: dict[int, Job] = {}
        latest_attempts: dict[int, JobAttempt] = {}
        latest_summaries: dict[int, DocumentTagSummary] = {}
        tool_counts: dict[int, dict[str, object]] = {}

        if doc_ids:
            jobs = session.exec(
                select(Job)
                .where(Job.document_id.in_(doc_ids))
                .order_by(Job.created_at.desc())
            ).all()
            for job in jobs:
                if job.document_id not in latest_jobs:
                    latest_jobs[job.document_id] = job

            job_ids = [job.id for job in latest_jobs.values()]
            if job_ids:
                attempts = session.exec(
                    select(JobAttempt)
                    .where(JobAttempt.job_id.in_(job_ids))
                    .order_by(JobAttempt.attempt_no.desc())
                ).all()
                for attempt in attempts:
                    if attempt.job_id not in latest_attempts:
                        latest_attempts[attempt.job_id] = attempt

            summaries = session.exec(
                select(DocumentTagSummary)
                .where(DocumentTagSummary.document_id.in_(doc_ids))
                .order_by(DocumentTagSummary.created_at.desc())
            ).all()
            for summary in summaries:
                if summary.document_id not in latest_summaries:
                    latest_summaries[summary.document_id] = summary

            if SHOW_TOOL_BADGES:
                job_ids = [job.id for job in latest_jobs.values()]
                if job_ids:
                    claims = session.exec(
                        select(DocumentClaim).where(DocumentClaim.job_id.in_(job_ids))
                    ).all()
                    risk_flags = session.exec(
                        select(DocumentRiskFlag).where(
                            DocumentRiskFlag.job_id.in_(job_ids)
                        )
                    ).all()
                    for claim in claims:
                        bucket = tool_counts.setdefault(
                            claim.job_id, {"citations": 0, "risk": {"high": 0, "medium": 0, "low": 0}}
                        )
                        if claim.quote_text:
                            bucket["citations"] += 1
                    for flag in risk_flags:
                        bucket = tool_counts.setdefault(
                            flag.job_id, {"citations": 0, "risk": {"high": 0, "medium": 0, "low": 0}}
                        )
                        severity = flag.severity or "low"
                        if severity not in bucket["risk"]:
                            bucket["risk"][severity] = 0
                        bucket["risk"][severity] += 1

        filtered_docs = []
        for doc in docs:
            job = latest_jobs.get(doc.id)
            audience_code = (
                job.audience or job.selected_audience or "auto" if job else "auto"
            )
            if selected_audience != "all" and audience_code != selected_audience:
                continue
            summary_row = latest_summaries.get(doc.id)
            if domain_filter and (not summary_row or summary_row.domain != domain_filter):
                continue
            attempt = latest_attempts.get(job.id) if job else None
            filtered_docs.append(
                {
                    "document_id": doc.id,
                    "job_id": job.id if job else None,
                    "audience": audience_code,
                    "status": job.status if job else "pending",
                    "snippet": (doc.content or "")[:200],
                    "summary": attempt.generated_one_line_summary if attempt else "",
                    "tool_summary": tool_counts.get(job.id) if job else None,
                }
            )

        start = (page - 1) * limit
        end = start + limit
        documents = filtered_docs[start:end]
        has_next = end < len(filtered_docs)

    return templates.TemplateResponse(
        request,
        "documents.html",
        {
            "audiences": AUDIENCES,
            "audience_display": audience_display,
            "selected_audience": selected_audience,
            "route_counts": route_counts,
            "documents": documents,
            "page": page,
            "has_prev": page > 1,
            "has_next": has_next,
            "query": search,
            "domain_filter": domain_filter,
            "show_tool_badges": SHOW_TOOL_BADGES,
        },
    )


@app.post("/web/documents", response_class=HTMLResponse)
def create_document(
    request: Request,
    input_text: str = Form(""),
    input_url: str = Form(""),
    use_sample: Optional[str] = Form(None),
    audience: str = Form("auto"),
    upload_file: Optional[UploadFile] = File(None),
) -> HTMLResponse:
    _check_rate_limit(request)
    content = ""
    source_type = ""
    input_text = (input_text or "").strip()
    input_url = (input_url or "").strip()

    if input_url and len(input_url) > MAX_URL_CHARS:
        context = _build_home_context(audience)
        return templates.TemplateResponse(
            request,
            "home.html",
            {
                **context,
                "sample_text": SAMPLE_TEXT,
                "error": "URL is too long.",
            },
            status_code=400,
        )

    if use_sample:
        content = SAMPLE_TEXT
        source_type = "sample"
    elif upload_file and upload_file.filename:
        fname = upload_file.filename.lower()
        if not (fname.endswith(".pdf") or fname.endswith(".docx")):
            context = _build_home_context(audience)
            return templates.TemplateResponse(
                request,
                "home.html",
                {
                    **context,
                    "sample_text": SAMPLE_TEXT,
                    "error": "Only PDF and Word (.docx) files are supported.",
                },
                status_code=400,
            )
        file_bytes = upload_file.file.read(MAX_UPLOAD_BYTES + 1)
        if len(file_bytes) > MAX_UPLOAD_BYTES:
            context = _build_home_context(audience)
            return templates.TemplateResponse(
                request,
                "home.html",
                {
                    **context,
                    "sample_text": SAMPLE_TEXT,
                    "error": "File is too large (max 10 MB).",
                },
                status_code=400,
            )
        content = _extract_file_text(file_bytes, upload_file.filename)
        source_type = "upload_pdf" if fname.endswith(".pdf") else "upload_docx"
    elif input_url:
        content = fetch_url_text(input_url)
        source_type = "url"
    elif input_text:
        content = input_text
        source_type = "text"

    if not content:
        error_message = "Provide text, a URL, upload a PDF/Word file, or choose sample content."
        if input_url:
            error_message = "URL fetch failed (blocked or empty content)."
        elif upload_file and upload_file.filename:
            error_message = "Could not extract text from the uploaded file. Is it a text-based (not scanned) PDF or valid .docx?"
        context = _build_home_context(audience)
        return templates.TemplateResponse(
            request,
            "home.html",
            {
                **context,
                "sample_text": SAMPLE_TEXT,
                "error": error_message,
            },
            status_code=400,
        )

    if audience not in AUDIENCES:
        raise HTTPException(status_code=400, detail="Invalid audience selection.")

    max_words = None
    if audience != "auto":
        max_words = get_audience_profile(audience).get("default_max_words")

    if len(content) > MAX_DOC_CHARS:
        content = content[:MAX_DOC_CHARS]

    doc = Document(content=content, source_type=source_type, source_url=input_url if source_type == "url" else None)
    job = Job(
        document_id=0,
        selected_audience=audience,
        audience=None if audience == "auto" else audience,
        max_words=max_words,
    )

    with get_session() as session:
        with session.begin():
            session.add(doc)
            session.flush()
            job.document_id = doc.id
            session.add(job)
        session.refresh(job)

    return RedirectResponse(url=f"/web/jobs/{job.id}", status_code=303)


@app.post("/web/jobs/{job_id}/run", response_class=HTMLResponse)
def run_job(request: Request, job_id: int) -> HTMLResponse:
    _check_rate_limit(request)
    with get_session() as session:
        job = session.get(Job, job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Job not found.")
        if job.status == "completed":
            return RedirectResponse(url=f"/web/jobs/{job_id}?analysis=1", status_code=303)
        if job.status == "running":
            return RedirectResponse(url=f"/web/jobs/{job_id}", status_code=303)

        # Atomically claim the job by setting status='running' only when it is still
        # in a runnable state.  Concurrent requests will see rowcount=0 and redirect.
        result = session.execute(
            text(
                "UPDATE job SET status = 'running'"
                " WHERE id = :id AND status IN ('pending', 'failed')"
            ).bindparams(id=job_id)
        )
        session.commit()
        if result.rowcount == 0:
            session.refresh(job)
            if job.status == "completed":
                return RedirectResponse(url=f"/web/jobs/{job_id}?analysis=1", status_code=303)
            return RedirectResponse(url=f"/web/jobs/{job_id}", status_code=303)

        session.refresh(job)
        try:
            run_job_pipeline(session, job)
        except Exception:
            logger.exception("Pipeline failed for job %d — job marked failed, redirecting.", job_id)
    return RedirectResponse(url=f"/web/jobs/{job_id}?analysis=1", status_code=303)


@app.get("/web/jobs/{job_id}", response_class=HTMLResponse)
def job_detail(request: Request, job_id: int) -> HTMLResponse:
    with get_session() as session:
        job = session.get(Job, job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Job not found.")
        doc = session.get(Document, job.document_id)

        attempts = session.exec(
            select(JobAttempt)
            .where(JobAttempt.job_id == job.id)
            .order_by(JobAttempt.attempt_no)
        ).all()

        tags = session.exec(
            select(Tag)
            .join(DocumentTag, Tag.id == DocumentTag.tag_id)
            .where(DocumentTag.document_id == job.document_id)
        ).all()

        clues = session.exec(
            select(DocumentClue).where(DocumentClue.document_id == job.document_id)
        ).all()

        raw_claims = session.exec(
            select(DocumentClaim)
            .where(DocumentClaim.document_id == job.document_id)
            .where(DocumentClaim.job_id == job.id)
            .order_by(DocumentClaim.created_at.desc())
        ).all()

        risk_flags = session.exec(
            select(DocumentRiskFlag)
            .where(DocumentRiskFlag.document_id == job.document_id)
            .where(DocumentRiskFlag.job_id == job.id)
            .order_by(DocumentRiskFlag.created_at.desc())
        ).all()

        # Related documents via tag Jaccard similarity
        summary_row = session.exec(
            select(DocumentTagSummary)
            .where(DocumentTagSummary.document_id == job.document_id)
            .order_by(DocumentTagSummary.created_at.desc())
        ).first()
        canonical_tags = parse_summary_tags(summary_row) if summary_row else []
        related_docs = _build_related_docs(session, job.document_id, canonical_tags, limit=8)

    audience_code = job.audience or job.selected_audience or "auto"
    audience_label = _audience_display(audience_code)
    routing_candidates = _safe_json_list(getattr(job, "routing_candidates_json", "[]"))
    routing_candidates_display = [
        _audience_display(aud) for aud in routing_candidates if isinstance(aud, str)
    ]
    routing_reasons = _safe_json_list(getattr(job, "routing_reasons_json", "[]"))
    routing_confidence_pct = None
    if getattr(job, "routing_confidence", None) is not None:
        routing_confidence_pct = int(round(float(job.routing_confidence) * 100))
    routing_source = getattr(job, "routing_source", None)

    ml_explanation = None
    if routing_source in ("ml", "ml+llm_fallback") and doc and doc.content:
        if _ml_router.load():
            ml_explanation = _ml_router.explain(doc.content, top_n=8)

    cross_functional_detail = None
    if audience_label == "Cross-Functional":
        if routing_candidates_display:
            cross_functional_detail = " + ".join(routing_candidates_display)
        else:
            cross_functional_detail = "Mixed stakeholders (auto-routed)"

    attempt_views = []
    final_summary = None
    for attempt in attempts:
        evaluator_data = _safe_json(attempt.evaluator_json)
        fail_reasons = evaluator_data.get("fail_reasons", [])
        attempt_views.append(
            {
                "attempt_id": attempt.id,
                "attempt_no": attempt.attempt_no,
                "passed": attempt.passed,
                "summary_preview": (attempt.generated_one_line_summary or "")[:200],
                "fail_reasons": fail_reasons,
            }
        )
        if attempt.passed:
            mind_map = (attempt.generated_mindmap or "").strip()
            final_summary = {
                "one_line_summary": attempt.generated_one_line_summary,
                "decision_bullets": _safe_json_list(attempt.generated_bullets_json),
                "mind_map": mind_map,
            }
            if final_summary["decision_bullets"] and not final_summary["mind_map"]:
                final_summary["mind_map"] = _fallback_mind_map_from_bullets(
                    final_summary["decision_bullets"]
                )

    claim_views = _build_claim_views(doc.content if doc else None, raw_claims)
    strong_count = len([c for c in claim_views if c["confidence"] >= 0.8])
    medium_count = len([c for c in claim_views if 0.6 <= c["confidence"] < 0.8])
    weak_count = len([c for c in claim_views if c["confidence"] < 0.6])
    supported_count = len([c for c in claim_views if c["quote_text"]])

    return templates.TemplateResponse(
        request,
        "job_detail.html",
        {
            "job": job,
            "audience_label": audience_label,
            "routing_candidates": routing_candidates_display,
            "routing_reasons": routing_reasons,
            "routing_confidence_pct": routing_confidence_pct,
            "cross_functional_detail": cross_functional_detail,
            "job_ref": f"JD-{job.id:06d}",
            "document": doc,
            "source_url": doc.source_url if doc else None,
            "attempts": attempt_views,
            "tags": tags,
            "clues": clues,
            "final_summary": final_summary,
            "auto_analysis": (
                request.query_params.get("analysis") == "1" and final_summary is not None
            ),
            "claims": claim_views,
            "risk_flags": risk_flags,
            "support_warning": supported_count < 3,
            "citation_summary": {
                "total": len(claim_views),
                "supported": supported_count,
                "strong": strong_count,
                "medium": medium_count,
                "weak": weak_count,
            },
            "related_docs": related_docs,
            "ml_explanation": ml_explanation,
        },
    )


@app.get("/web/documents/{doc_id}", response_class=HTMLResponse)
def document_detail(request: Request, doc_id: int) -> HTMLResponse:
    with get_session() as session:
        doc = session.get(Document, doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found.")

        job = session.exec(
            select(Job)
            .where(Job.document_id == doc.id)
            .order_by(Job.created_at.desc())
        ).first()

        summary_row = session.exec(
            select(DocumentTagSummary)
            .where(DocumentTagSummary.document_id == doc.id)
            .order_by(DocumentTagSummary.created_at.desc())
        ).first()
        canonical_tags = parse_summary_tags(summary_row) if summary_row else []
        domain_label = summary_row.domain if summary_row else None

        related_docs = _build_related_docs(session, doc.id, canonical_tags, limit=10)

        final_summary = None
        if job:
            attempt = session.exec(
                select(JobAttempt)
                .where(JobAttempt.job_id == job.id)
                .where(JobAttempt.passed == True)  # noqa: E712
                .order_by(JobAttempt.attempt_no.desc())
            ).first()
            if attempt:
                mind_map = (attempt.generated_mindmap or "").strip()
                final_summary = {
                    "one_line_summary": attempt.generated_one_line_summary,
                    "decision_bullets": _safe_json_list(attempt.generated_bullets_json),
                    "mind_map": mind_map,
                }
                if final_summary["decision_bullets"] and not final_summary["mind_map"]:
                    final_summary["mind_map"] = _fallback_mind_map_from_bullets(
                        final_summary["decision_bullets"]
                    )

        raw_claims = []
        risk_flags = []
        support_warning = False
        if job:
            raw_claims = session.exec(
                select(DocumentClaim)
                .where(DocumentClaim.document_id == doc.id)
                .where(DocumentClaim.job_id == job.id)
                .order_by(DocumentClaim.created_at.desc())
            ).all()
            risk_flags = session.exec(
                select(DocumentRiskFlag)
                .where(DocumentRiskFlag.document_id == doc.id)
                .where(DocumentRiskFlag.job_id == job.id)
                .order_by(DocumentRiskFlag.created_at.desc())
            ).all()
            support_warning = len([c for c in raw_claims if c.quote_text]) < 3

        claim_views = _build_claim_views(doc.content, raw_claims)
        tags = session.exec(
            select(Tag)
            .join(DocumentTag, Tag.id == DocumentTag.tag_id)
            .where(DocumentTag.document_id == doc.id)
        ).all()

        clues = session.exec(
            select(DocumentClue).where(DocumentClue.document_id == doc.id)
        ).all()

    audience_label = None
    if job:
        audience_label = _audience_display(job.audience or job.selected_audience or "auto")
    routing_candidates_display = []
    routing_reasons = []
    routing_confidence_pct = None
    cross_functional_detail = None
    if job:
        routing_candidates = _safe_json_list(getattr(job, "routing_candidates_json", "[]"))
        routing_candidates_display = [
            _audience_display(aud) for aud in routing_candidates if isinstance(aud, str)
        ]
        routing_reasons = _safe_json_list(getattr(job, "routing_reasons_json", "[]"))
        if getattr(job, "routing_confidence", None) is not None:
            routing_confidence_pct = int(round(float(job.routing_confidence) * 100))
        if audience_label == "Cross-Functional":
            if routing_candidates_display:
                cross_functional_detail = " + ".join(routing_candidates_display)
            else:
                cross_functional_detail = "Mixed stakeholders (auto-routed)"

    return templates.TemplateResponse(
        request,
        "document_detail.html",
        {
            "document": doc,
            "job": job,
            "audience_label": audience_label,
            "routing_candidates": routing_candidates_display,
            "routing_reasons": routing_reasons,
            "routing_confidence_pct": routing_confidence_pct,
            "cross_functional_detail": cross_functional_detail,
            "final_summary": final_summary,
            "tags": tags,
            "clues": clues,
            "domain_label": domain_label,
            "canonical_tags": canonical_tags,
            "related_docs": related_docs,
            "claims": claim_views,
            "risk_flags": risk_flags,
            "support_warning": support_warning,
        },
    )


@app.get("/web/attempts/{attempt_id}", response_class=HTMLResponse)
def attempt_detail(request: Request, attempt_id: int) -> HTMLResponse:
    with get_session() as session:
        attempt = session.get(JobAttempt, attempt_id)
        if not attempt:
            raise HTTPException(status_code=404, detail="Attempt not found.")
        job = session.get(Job, attempt.job_id)
        document = session.get(Document, job.document_id) if job else None

    return templates.TemplateResponse(
        request,
        "attempt_detail.html",
        {
            "attempt": attempt,
            "job": job,
            "document": document,
            "routing_reasons": _safe_json_list(
                getattr(job, "routing_reasons_json", "[]")
            )
            if job
            else [],
            "generated_tags": _safe_json_list(attempt.generated_tags_json),
            "generated_clues": _safe_json_list(attempt.generated_clues_json),
            "generated_bullets": _safe_json_list(attempt.generated_bullets_json),
            "evaluator_data": _safe_json(attempt.evaluator_json),
        },
    )


def _extract_file_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from an uploaded PDF or Word (.docx) file."""
    name_lower = (filename or "").lower()
    try:
        if name_lower.endswith(".pdf"):
            import pypdf  # imported here so the app still starts without the package
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            parts = [page.extract_text() or "" for page in reader.pages]
            return " ".join(parts)
        if name_lower.endswith(".docx"):
            import docx  # python-docx
            doc = docx.Document(io.BytesIO(file_bytes))
            return " ".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        logger.exception("File text extraction failed for %s", filename)
    return ""


def fetch_url_text(url: str) -> str:
    """Fetch plain text from a URL.

    Hardening notes:
    - Content-Type is checked before the body is parsed; non-HTML/text responses are rejected.
    - Raw bytes are capped at MAX_FETCH_BYTES before BeautifulSoup sees them.
    - Timeout is capped at _URL_FETCH_TIMEOUT per hop to keep demo runs responsive.
    - Known limitation: DNS rebinding (TOCTOU) is not fully mitigated. _is_safe_url()
      resolves DNS once; httpx resolves again independently. A full fix requires a custom
      httpx transport that reuses the validated IP — left as a production hardening item.
    """
    url = (url or "").strip()
    if not url or len(url) > MAX_URL_CHARS:
        return ""

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0 Safari/537.36"
        )
    }

    for _ in range(MAX_URL_REDIRECTS + 1):
        if not _is_safe_url(url):
            logger.warning("Blocked unsafe URL fetch: %s", url[:200])
            return ""
        try:
            with httpx.stream(
                "GET",
                url,
                timeout=_URL_FETCH_TIMEOUT,
                follow_redirects=False,
                headers=headers,
            ) as response:
                if response.is_redirect:
                    location = response.headers.get("location")
                    if not location:
                        return ""
                    url = urljoin(str(response.url), location)
                    continue
                response.raise_for_status()

                # Reject non-text content types before reading the body.
                content_type = response.headers.get("content-type", "").lower()
                if not any(content_type.startswith(p) for p in _ALLOWED_CONTENT_TYPE_PREFIXES):
                    logger.warning(
                        "Blocked URL fetch — unsupported content-type %r: %s",
                        content_type[:80],
                        url[:200],
                    )
                    return ""

                # Stream body in chunks — stop reading at MAX_FETCH_BYTES so memory
                # is bounded even if the server sends an unbounded response.
                chunks: list[bytes] = []
                total = 0
                for chunk in response.iter_bytes(chunk_size=8192):
                    chunks.append(chunk)
                    total += len(chunk)
                    if total >= MAX_FETCH_BYTES:
                        logger.warning(
                            "URL response truncated at %d bytes: %s",
                            MAX_FETCH_BYTES,
                            url[:200],
                        )
                        break
                raw = b"".join(chunks)[:MAX_FETCH_BYTES]
        except Exception:
            logger.exception("URL fetch failed.")
            return ""

        html_text = raw.decode("utf-8", errors="replace")
        soup = BeautifulSoup(html_text, "lxml")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = " ".join(soup.stripped_strings)
        return text[:MAX_DOC_CHARS]

    return ""


def _safe_json(payload: str) -> dict:
    try:
        return json.loads(payload) if payload else {}
    except json.JSONDecodeError:
        return {}


def _safe_json_list(payload: str) -> list:
    data = _safe_json(payload)
    if isinstance(data, list):
        return data
    return []


def _fallback_mind_map_from_bullets(bullets: list[str]) -> str:
    lines = ["mindmap", "  root((Summary))"]
    for bullet in bullets[:5]:
        title = str(bullet).split(":", 1)[0].strip()
        if not title:
            title = str(bullet).strip()[:40]
        if title:
            lines.append(f"    {title}")
    return "\n".join(lines)
